# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "out_delivery" / "classification_readme.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, col_widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(row.cells[idx])

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, 11)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    if level == 1:
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(8)
        size, color = 14, "1F4D78"
    elif level == 2:
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        size, color = 12, "1F4D78"
    else:
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        size, color = 11, "1F4D78"
    if ". " in text:
        prefix, rest = text.split(". ", 1)
        run = paragraph.add_run(prefix + ". ")
        set_run_font(run, size=size, bold=True, color=color)
        run = paragraph.add_run(rest)
        set_run_font(run, size=size, bold=False, color=color)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=False, color=color)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style=None)
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run("• ")
    set_run_font(run, 11)
    run = paragraph.add_run(text)
    set_run_font(run, 11)
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                style_paragraph(paragraph, after=2, line=1.15)
                for run in paragraph.runs:
                    set_run_font(run, 10)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("广州鸡类餐饮 POI 分类口径说明")
    set_run_font(run, 20, bold=True, color="0B2545")

    subtitle = doc.add_paragraph("适用文件：spike/out_delivery/poi_clean.csv 与 poi_clean.geojson")
    style_paragraph(subtitle, after=10)

    add_heading(doc, "1. 口径定位", 1)
    add_body(
        doc,
        "本说明用于统一第四章“广州样本”中鸡类餐饮 POI 的分类口径。分类对象不是整家店的菜系，"
        "而是该 POI 的店名、推荐菜或标签中被识别出来的鸡类菜品信号。"
    )
    add_bullet(doc, "研究目标：观察鸡文化在广州现代餐饮空间中的可见度、隐性渗透和传统/现代分化。")
    add_bullet(doc, "主数据：清洗后的高德 POI，正式交付目录为 spike/out_delivery/。")
    add_bullet(doc, "解释边界：本数据不等同于居民实际吃鸡消费量、订单量或营业额。")

    add_heading(doc, "2. 数据来源与坐标系", 1)
    add_body(
        doc,
        "本 README 固化的是 POI 清洗分类口径。数据来源分为参与计算的高德 POI 原始数据、"
        "用于故事背景的公开资料整理表，以及最终交付给制图环节的清洗结果。"
    )
    add_table(
        doc,
        ["数据项", "来源与用途", "参与计算"],
        [
            ["鸡类候选 POI", "高德地图 POI 原始抓取结果；用于清洗、去重和分类。", "是"],
            ["全量餐饮 POI", "高德地图餐饮服务全量样本；作为各区渗透率分母。", "是"],
            ["背景资料源表", "公开资料整理表；用于第四章故事背景和外部佐证。", "否"],
            ["清洗后 POI", "正式交付给 GeoScene 点图层、分类结构图和空间分析。", "结果"],
            ["分区渗透率表", "由鸡类 POI 分子与全量餐饮 POI 分母计算生成。", "结果"],
        ],
        [1.45, 4.2, 0.85],
    )
    add_bullet(doc, "鸡类候选 POI 原始文件：spike/data/poi_raw_chicken_keyword.jsonl。")
    add_bullet(doc, "全量餐饮 POI 原始文件：spike/data/poi_raw_all_food.jsonl。")
    add_bullet(doc, "背景资料整理表：spike/data/background_context_sources.csv，包含统计年鉴、政府新闻、行业报告、协会报告和媒体报道，仅用于故事背景和引用线索，不参与 POI 分类或渗透率计算。")
    add_bullet(doc, "正式交付文件：spike/out_delivery/poi_clean.csv、poi_clean.geojson 与 penetration_by_district.csv。")
    add_bullet(doc, "原始高德 POI 的 location 字段为 GCJ-02 坐标。")
    add_bullet(doc, "当前 out_delivery 交付数据由 python3 spike/build_delivery_data.py 生成，未使用 --keep-gcj02，因此 lng/lat 输出为 WGS-84。")
    add_bullet(doc, "如果后续使用高德或其他 GCJ-02 底图，应另行用 --keep-gcj02 导出一版，不能与当前 WGS-84 结果混用。")

    add_heading(doc, "3. 分类说明", 1)
    add_table(
        doc,
        ["标签", "含义", "典型信号", "解释方式"],
        [
            ["traditional", "广东本地做法、地方鸡种或本地饭式", "白切鸡、豉油鸡、盐焗鸡、走地鸡、窑鸡、煲仔饭", "传统鸡文化在城市餐饮中的延续"],
            ["modern", "外地做法、现代快餐化或连锁化鸡类产品", "炸鸡、鸡排、鸡块、鸡腿堡、鸡肉披萨、竹筒饭、普通烧鸡", "快餐化和外来鸡类消费场景扩张"],
            ["other_chicken", "有鸡类信号，但做法或地域归属不清", "鸡肉、鸡饭、鸡煲、鸡火锅", "只说明能识别出鸡，不强行归入传统或现代"],
            ["non_chicken", "没有足够鸡菜证据", "只有海鲜、烧鹅、火锅、酒家、茶餐厅等业态词", "不进入鸡类菜品结构分析"],
            ["excluded", "含“鸡”但不是鸡肉餐饮", "田鸡、鸡尾酒、鸡蛋仔、鸡精", "清洗阶段剔除"],
        ],
        [1.25, 1.65, 2.1, 1.5],
    )

    add_heading(doc, "4. 分类流程", 1)
    add_body(doc, "同一条 POI 可能同时包含多个菜品词。当前规则按以下优先级判定：")
    for item in [
        "假阳性词优先剔除，例如田鸡、鸡尾酒、鸡蛋仔。",
        "强现代信号优先判为 modern，例如麦辣鸡、板烧鸡、鸡块、鸡排、鸡腿堡、鸡肉披萨、竹筒饭。",
        "“烧鸡”默认判为 modern；若与客家、粤菜、广府、顺德、潮汕、清远、湛江、荔枝木、走地、土鸡等本地语境共同出现，则判为 traditional。",
        "广东本地鸡菜、地方鸡种和本地饭式判为 traditional。",
        "明确外地、现代或连锁化鸡菜判为 modern。",
        "只含泛鸡信号且无法判断做法时判为 other_chicken。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 关键争议项裁决", 1)
    add_table(
        doc,
        ["对象", "最终口径", "理由"],
        [
            ["麦当劳、肯德基、赛百味、汉堡王、华莱士", "modern", "鸡块、鸡排、鸡腿堡、板烧鸡等属于现代快餐化鸡类产品。"],
            ["尊宝、达美乐等披萨品牌", "modern", "鸡肉披萨、奥尔良烤鸡、照烧鸡肉等属于现代或西式餐饮语境。"],
            ["煲仔饭", "traditional", "按广东本地饭式处理；即使 tag 不完整，也保留生活经验口径。"],
            ["竹筒饭", "modern", "按土家/傣族等非广东本地传统食物语境处理；若与农家菜同现，以竹筒饭优先。"],
            ["普通烧鸡", "modern", "无地域修饰时按现代或外来烧鸡处理。"],
            ["客家烧鸡、荔枝木烧鸡、窑鸡", "traditional", "带广东本地或客家语境，作为本地鸡菜信号。"],
            ["卤鸡腿", "modern", "按非广东本地鸡菜处理；但若同一 tag 同时有白切鸡等传统词，当前可能被传统词覆盖。"],
            ["火锅", "other_chicken 或 non_chicken", "火锅是业态而不是鸡菜；只有鸡肉/鸡火锅信号时进入 other_chicken。"],
            ["酒家、大排档、海鲜城、茶餐厅、农庄", "不单独决定分类", "这些是餐厅类型，最终分类取决于其中出现的鸡类菜品信号。"],
            ["砂锅粥、烧鹅、烧烤", "不单独作为鸡类信号", "除非同一条记录中出现白切鸡、土鸡、鸡肉等鸡菜证据。"],
        ],
        [1.45, 1.25, 3.8],
    )

    add_heading(doc, "6. 图表使用建议", 1)
    add_bullet(doc, "做“传统与现代”结构图时，只比较 traditional 与 modern，other_chicken 单独作为无法归属类展示。")
    add_bullet(doc, "做“鸡类渗透率”时，使用 signboard_pct 与 serves_pct；不要把它写成真实消费量。")
    add_bullet(doc, "解读酒家、海鲜城、大排档等综合餐厅时，应说明分类来自菜单鸡菜信号，不代表整家店只经营该类菜。")

    add_heading(doc, "7. 当前版本数据摘要", 1)
    add_table(
        doc,
        ["指标", "当前值"],
        [
            ["clean_rows", "14,091"],
            ["traditional", "3,570"],
            ["modern", "4,560"],
            ["other_chicken", "2,998"],
            ["non_chicken", "2,963"],
            ["contains_chicken_rows", "11,128"],
            ["signboard_rows", "7,405"],
            ["recommended_dish_rows", "3,723"],
        ],
        [2.7, 3.8],
    )

    add_heading(doc, "8. 后续改进口径", 1)
    add_body(
        doc,
        "下一步建议在 CSV 中新增 matched_terms、matched_rule、confidence 字段。这样看到“海鲜城 = traditional”时，"
        "可以直接知道它是因为白切鸡、走地鸡等菜品词命中，而不是因为海鲜城这个业态本身。"
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
